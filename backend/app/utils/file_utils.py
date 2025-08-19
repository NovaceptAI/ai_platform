import os
import docx
try:
	import PyPDF2
except Exception:
	PyPDF2 = None
try:
	from pydub import AudioSegment
except Exception:
	AudioSegment = None
try:
	import moviepy.editor as mp
except Exception:
	mp = None
from app.services.lemonfox_service import transcribe_audio
from app.services.whisper_service import translate_audio_with_azure_whisper
from transformers import GPT2TokenizerFast
# transcriber = AzureWhisperTranscriber()

tokenizer = GPT2TokenizerFast.from_pretrained("gpt2")
TOKEN_LIMIT = 1500  # You can tweak this

def detect_file_type(file_path):
	"""Detects the file type based on its extension."""
	if file_path.endswith(('.docx', '.pdf', '.txt')):
		return "document"
	elif file_path.endswith(('.mp3', '.wav', '.m4a')):
		return "audio"
	elif file_path.endswith(('.mp4', '.avi', '.mov')):
		return "video"
	else:
		raise ValueError("Unsupported file format")

def extract_text_from_document(file_path):
	"""Returns a list of chunks: pages for PDFs, token-wise for DOCX/TXT."""
	if file_path.endswith('.pdf'):
		if PyPDF2 is None:
			raise RuntimeError("PyPDF2 not available")
		with open(file_path, 'rb') as file:
			reader = PyPDF2.PdfReader(file)
			return [page.extract_text() or "" for page in reader.pages]

	elif file_path.endswith('.docx'):
		doc = docx.Document(file_path)
		text = '\n'.join([para.text for para in doc.paragraphs])
		return _chunk_by_token(text)

	elif file_path.endswith('.txt'):
		with open(file_path, 'r', encoding='utf-8') as f:
			text = f.read()
		return _chunk_by_token(text)

	else:
		raise ValueError("Unsupported document format")


def _chunk_by_token(text, token_limit=TOKEN_LIMIT):
	"""Splits text into token-sized chunks."""
	words = text.split()
	chunks, current_chunk = [], []

	for word in words:
		current_chunk.append(word)
		if len(tokenizer.encode(' '.join(current_chunk))) > token_limit:
			chunks.append(' '.join(current_chunk[:-1]))
			current_chunk = [word]
	
	if current_chunk:
		chunks.append(' '.join(current_chunk))
	
	return chunks


def extract_text_from_audio(file_path):
	"""Extracts and transcribes audio to text."""
	if translate_audio_with_azure_whisper is None:
		raise RuntimeError("Azure Whisper service not available")
	try:
		return translate_audio_with_azure_whisper(file_path)
	finally:
		if os.path.exists(file_path):
			os.remove(file_path)

def extract_text_from_video(file_path):
	"""Extracts audio from video and transcribes it to text."""
	if mp is None:
		raise RuntimeError("moviepy not available")
	temp_audio_path = "temp_audio.wav"
	video = mp.VideoFileClip(file_path)
	video.audio.write_audiofile(temp_audio_path)
	try:
		return translate_audio_with_azure_whisper(temp_audio_path)
	finally:
		if os.path.exists(temp_audio_path):
			os.remove(temp_audio_path)


def extract_text_by_pages(file_path):
	"""Extracts page-wise text (PDF), or chunked token-based segments (DOCX/TXT)."""
	if file_path.endswith('.pdf'):
		if PyPDF2 is None:
			raise RuntimeError("PyPDF2 not available")
		with open(file_path, 'rb') as file:
			reader = PyPDF2.PdfReader(file)
			return [page.extract_text() or "" for page in reader.pages]

	elif file_path.endswith('.docx'):
		doc = docx.Document(file_path)
		# Chunk every N paragraphs as a 'page'
		paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
		chunk_size = 5
		return ['\n'.join(paragraphs[i:i+chunk_size]) for i in range(0, len(paragraphs), chunk_size)]

	elif file_path.endswith('.txt'):
		with open(file_path, 'r', encoding='utf-8') as f:
			text = f.read()
		tokens = text.split()
		chunk_size = 300
		return [' '.join(tokens[i:i+chunk_size]) for i in range(0, len(tokens), chunk_size)]

	else:
		raise ValueError("Unsupported document format")